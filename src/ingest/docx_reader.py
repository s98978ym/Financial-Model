"""DOCX document text extraction with estimated page numbers.

Uses python-docx to read paragraphs and tables.  Because the DOCX format
does not embed hard page numbers in a way that is trivially accessible,
this module uses two heuristics:

1. Explicit page-break detection (``<w:br w:type="page"/>`` runs).
2. A configurable line-count estimator as a secondary fallback.

Tables are extracted as structured data alongside the surrounding text.
"""
import logging
import re
from pathlib import Path
from typing import List, Optional

from .base import DocumentContent, PageContent

logger = logging.getLogger(__name__)

# Approximate number of paragraph elements that fit on one printed page.
# Used only when no explicit page breaks are found.
_PARAGRAPHS_PER_PAGE_ESTIMATE = 40


def _has_page_break(paragraph) -> bool:
    """Return True if *paragraph* contains an explicit page break run."""
    # python-docx exposes the underlying XML; look for <w:br w:type="page"/>
    try:
        from lxml import etree  # type: ignore

        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        for br in paragraph._element.findall(".//w:br", ns):
            if br.get(f"{{{ns['w']}}}type") == "page":
                return True
    except Exception:
        # If lxml is unavailable or the XML walk fails, fall through.
        pass
    return False


def _extract_table_data(table) -> List[List[str]]:
    """Convert a python-docx Table object into a list-of-rows."""
    rows: List[List[str]] = []
    try:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
    except Exception as exc:
        logger.warning("Failed to extract table: %s", exc)
    return rows


def extract_docx(file_path: str) -> DocumentContent:
    """Extract text and tables from a DOCX file.

    Parameters
    ----------
    file_path : str
        Path to the DOCX file.

    Returns
    -------
    DocumentContent
        Extracted document with per-page text and tables.

    Raises
    ------
    FileNotFoundError
        If *file_path* does not exist.
    ImportError
        If python-docx is not installed.
    RuntimeError
        If the file cannot be read.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {file_path}")
    if not path.is_file():
        raise ValueError(f"Path is not a file: {file_path}")

    try:
        import docx  # type: ignore  (python-docx)
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX extraction. "
            "Install it with `pip install python-docx`."
        )

    try:
        doc = docx.Document(file_path)
    except Exception as exc:
        raise RuntimeError(
            f"Failed to open DOCX file: {file_path} — {exc}"
        ) from exc

    # ------------------------------------------------------------------
    # Extract core properties as metadata
    # ------------------------------------------------------------------
    metadata: dict = {}
    try:
        cp = doc.core_properties
        if cp.title:
            metadata["Title"] = cp.title
        if cp.author:
            metadata["Author"] = cp.author
        if cp.subject:
            metadata["Subject"] = cp.subject
        if cp.created:
            metadata["Created"] = str(cp.created)
        if cp.modified:
            metadata["Modified"] = str(cp.modified)
    except Exception:
        pass

    # ------------------------------------------------------------------
    # Walk the document body in order.
    #
    # The body can contain paragraphs and tables interleaved.  We iterate
    # over the underlying XML children to preserve their document order.
    # ------------------------------------------------------------------
    from docx.oxml.ns import qn  # type: ignore

    body = doc.element.body
    para_index = 0
    table_index = 0

    # Gather all paragraph and table python-docx objects for easy lookup.
    all_paragraphs = doc.paragraphs
    all_tables = doc.tables

    # Track page construction.
    current_page_number = 1
    current_texts: List[str] = []
    current_tables: List[List[List[str]]] = []
    paragraph_count_in_page = 0

    pages: List[PageContent] = []

    def _flush_page() -> None:
        """Append the accumulated content as a PageContent."""
        nonlocal current_texts, current_tables, current_page_number
        text = "\n".join(current_texts).strip()
        if text or current_tables:
            pages.append(
                PageContent(
                    page_number=current_page_number,
                    text=text,
                    tables=current_tables,
                    source_type="docx",
                )
            )
        current_page_number += 1
        current_texts = []
        current_tables = []

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            # Map XML element back to python-docx Paragraph
            if para_index < len(all_paragraphs):
                para = all_paragraphs[para_index]

                # Check for explicit page break *before* this paragraph's text.
                if _has_page_break(para) and (current_texts or current_tables):
                    _flush_page()
                    paragraph_count_in_page = 0

                para_text = para.text or ""
                if para_text.strip():
                    current_texts.append(para_text)

                paragraph_count_in_page += 1

                # Heuristic page break by paragraph count
                if paragraph_count_in_page >= _PARAGRAPHS_PER_PAGE_ESTIMATE:
                    _flush_page()
                    paragraph_count_in_page = 0

            para_index += 1

        elif tag == "tbl":
            if table_index < len(all_tables):
                table_data = _extract_table_data(all_tables[table_index])
                if table_data:
                    current_tables.append(table_data)
                    # Also add a text representation of the table for full-text search.
                    table_lines = []
                    for row in table_data:
                        table_lines.append(" | ".join(row))
                    current_texts.append("\n".join(table_lines))
            table_index += 1

    # Flush remaining content.
    if current_texts or current_tables:
        _flush_page()

    # Handle completely empty documents.
    total_pages = len(pages)
    if total_pages == 0:
        pages.append(
            PageContent(
                page_number=1,
                text="",
                tables=[],
                source_type="docx",
            )
        )
        total_pages = 1

    return DocumentContent(
        file_path=str(file_path),
        file_type="docx",
        pages=pages,
        total_pages=total_pages,
        metadata=metadata,
    )
